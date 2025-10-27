import pandas as pd
import os
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.section import WD_ORIENT, WD_ORIENTATION
from datetime import datetime

# Create Main KPIs
def master_table_summary(df):
    file_cnt = len(df['FileName'].unique())
    Column_cnt = df['ColumnCnt'].sum()
    record_Cnt = df['RecordCnt'].sum()/100000
    file_size = df['FileSize'].sum()/1000000
    work_date = df['WorkDate'].unique().max()
    return {
        "Total Files #": f"{file_cnt:,}", 
        "Total Columns #": f"{Column_cnt:,}", # Column_cnt,
        "Record #(백만)": f"{record_Cnt:,.0f}",   # record_Cnt, 
        "File Size(MB)": f"{file_size:,.0f}",  #file_size,
        "Analyzed Date" : work_date
    }

def Display_KPIs(df):
    st.markdown("####  분석 대상 파일(파일)들의 주요 지표 입니다.")   
    # CSS 스타일 정의
    st.markdown("""
        <style>
        div[data-testid="metric-container"] {
            background-color: #FFFFFF;
            border: 1px solid #E0E0E0;
            padding: 1rem;
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        div[data-testid="metric-container"] > label[data-testid="stMetricLabel"] {
            display: flex;
            justify-content: center;
            color: #404040;
            font-weight: 600;
        }
        div[data-testid="metric-container"] > div[data-testid="stMetricValue"] {
            display: flex;
            justify-content: center;
            font-weight: bold;
        }
        </style>
    """, unsafe_allow_html=True)

    # KPI 표시
    summary = master_table_summary(df)
    
    # 각 메트릭에 대한 색상 정의
    metric_colors = {
        "Total Files #": "#1f77b4",      # 파란색
        "Total Columns #": "#2ca02c",      # 초록색
        "Record #(백만)": "#ff7f0e",      # 주황색
        "File Size(MB)": "#d62728",       # 빨간색
        "Working Date": "#9467bd"         # 보라색
    }

    # 메트릭 표시
    cols = st.columns(len(summary))
    for col, (key, value) in zip(cols, summary.items()):
        color = metric_colors.get(key, "#0072B2")  # 기본 색상
        col.markdown(f"""
            <div style="text-align: center; padding: 1.2rem; background-color: #FFFFFF; border-radius: 10px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <div style="color: {color}; font-size: 2em; font-weight: bold;">{value}</div>
                <div style="color: #404040; font-size: 1em; margin-top: 0.5rem;">{key}</div>
            </div>
        """, unsafe_allow_html=True)

# -------------------------------------------------------------------------------
# 보고서 생성 함수
# -------------------------------------------------------------------------------
def generate_text_report(df, CURRENT_DIR_PATH, QDQM_ver):
    if df.empty:
        st.warning("데이터가 없습니다.")
        return

    # 보고서 내용 생성
    report_content = "\n"  
    report_content += """
══════════════════════════════════════════════════════════════
                                                             
               Data Quality Management Report                 
                                                              
                   QDQM Analyzer Report                      
                                                             
══════════════════════════════════════════════════════════════
"""
    report_content += "\n"*3   
    report_content += f"""
    \t작성자\t\t: QDQM Analyzer
    \t작성일시\t\t: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}
    \t버전\t\t: {QDQM_ver}
    """
    report_content += "\n"*3

    report_content += """
    \t본 보고서는 QDQM Analyzer에 의해 자동 생성된 보고서입니다.
    \t데이터 품질 관리를 위한 자료로 활용하시기 바랍니다.
    """
    report_content += "\n" + "=" * 60 + "\n\n"
    
    # 기본 정보 추가
    summary = master_table_summary(df)
    report_content += "[ Tables Information ]\n\n"
    for key, value in summary.items():
        report_content += f"\t{key}\t: {value}\n"

    unicode_files = df[df['UnicodeCnt'] > 0]['FileName'].tolist()
    unicode_file_count = len(unicode_files)
    report_content += f"\n\tUnicode 문자 포함 파일 수: {unicode_file_count} 개\n"
    
    broken_kor_files = df[df['BrokenKoreanCnt'] > 0]['FileName'].tolist()
    broken_kor_file_count = len(broken_kor_files)
    report_content += f"\t한글 미완성 문자 포함 파일 수: {broken_kor_file_count} 개\n"
    
    lt_lcl_files = df[df['BelowLCL'] > 0]['FileName'].tolist()
    lt_lcl_file_count = len(lt_lcl_files)
    report_content += f"\tLCL 보다 작은 값을 갖고있는 파일 수: {lt_lcl_file_count} 개\n"  
    
    gt_ucl_files = df[df['UpperUCL'] > 0]['FileName'].tolist()
    gt_ucl_file_count = len(gt_ucl_files)
    report_content += f"\tUCL 보다 큰 값을 갖고있는 파일 수: {gt_ucl_file_count} 개\n"   
    
      
    LCLUCL_Site = "https://qliksense.tistory.com/45"
    report_content += f"\n\tLCL, UCL에 관한 자세한 내용은 여기를 참조하세요: {LCLUCL_Site}\n"
    
    report_content += "\n" + "=" * 60 + "\n\n"

    # 파일별 상세 정보
    report_content += "[ Tables Detail Information ]\n"
    report_content += "\n" + "=" * 60 + "\n\n"
    for _, row in df.iterrows():
        report_content += f"[ {row['FileName']} ] Table Detail Information \n\n"
        report_content += f"\tColumns #\t\t: {row['ColumnCnt']}\n"
        report_content += f"\tFile Size\t\t: {row['FileSize']/1000:,.0f} KB\n"
        if len(str(int(row['RecordCnt']))) > 6:
            report_content += f"\tRow #\t\t: {row['RecordCnt']:,.0f}\n"
        else:
            report_content += f"\tRow #\t\t: {row['RecordCnt']:,.0f}\n"
        report_content += f"\tSampling #\t: {row['SamplingRows']:,.0f}\n"
        report_content += f"\tUnique Columns\t: {row['UniqueCnt']}\n"
        report_content += "\n\t--- Column # by Data Type ---\n"
        report_content += f"\tText Columns\t: {row['IsText']}\t\t"
        report_content += f"\tNull Columns\t: {row['IsNull']}\n"
        report_content += f"\tNumeric Columns\t: {row['IsNum']}\t\t"
        report_content += f"\tVarchar(Num)\t: {row['NumChar']}\n"    
        report_content += f"\tDate Columns\t: {row['Date']}\t\t"
        report_content += f"\tVarchar(Date)\t: {row['DateChar']}\n"
        report_content += f"\tVarchar(Time)\t: {row['TimeChar']}\t\t"
        report_content += f"\tTimeStamp Columns\t: {row['TimeStamp']}\n" 
        report_content += "\n\t--- Column # by Data Value ---\n"
        report_content += f"\tHas Alpha     \t: {row['HasAlpha']}\n"
        report_content += f"\tHas Only Alpha\t: {row['HasOnlyAlpha']}\t\t"
        report_content += f"\tOnly Alpha&Num\t: {row['HasOnlyAlphaNum']}\n"
        report_content += f"\tHas Kor    \t: {row['HasKor']}\n"
        report_content += f"\tHas Only Kor\t: {row['HasOnlyKor']}\t\t"
        report_content += f"\tBroken Kor\t: {row['BrokenKoreanCnt']}\n"
        report_content += f"\tHas Special\t: {row['SpecialCnt']}\t\t"
        report_content += f"\tHas Unicode\t: {row['UnicodeCnt']}\n"
        report_content += f"\tHas Chinese\t: {row['ChineseCnt']}\t\t"   
        report_content += f"\tHas Null   \t: {row['NullCnt']}\n"
        report_content += f"\tHas Blank  \t: {row['HasBlank']}\t\t"
        report_content += f"\tHas Dash  \t: {row['HasDash']}\n"
        report_content += f"\tHas Dot    \t: {row['HasDot']}\t\t"
        report_content += f"\tHas Bracket\t: {row['HasBracket']}\n"
        # report_content += f"Has At\t: {row['HasAt_Cnt']}\n"
        report_content += "\n\t--- Column # by Numeric Value & Statistics ---\n"
        report_content += f"\tHas Only Num\t: {row['HasOnlyNum']}\n"
        report_content += f"\tHas Minus   \t: {row['HasMinus']}\t\t"
        # report_content += f"\tHas Float  \t: {row['HasReal']}\n"
        report_content += f"\tLess Than LCL\t: {row['BelowLCL']}\t\t"
        report_content += f"\tGreater Than UCL\t: {row['UpperUCL']}\n"

        report_content += "\n\t--- Column # by Master Matching ---\n"
        report_content += f"\t적합        \t: {row['적합']}\t\t"
        report_content += f"\t양호        \t: {row['양호']}\n"
        report_content += f"\t고려        \t: {row['고려']}\t\t"
        report_content += f"\t부적합      \t: {row['부적합']}\n"
        report_content += f"\tNot Matched  \t: {row['X']}\t\t"
        report_content += f"\tMatching Ratio\t: {(row['적합']+row['양호'])/(row['ColumnCnt']):,.0%}\n"
        report_content += "\tMatching Ratio = (적합 + 양호) / 컬럼수\n"
        report_content += "-" * 100 + "\n\n"

    # # Unicode 문자가 포함된 파일 분석
    # unicode_files = df[df['HasUnicode_Cnt'] > 0]['FileName'].tolist()
    # unicode_file_count = len(unicode_files)
    
    # report_content += "\n[ Unicode 문자 포함 파일 분석 ]\n\n"
    # report_content += f"    Unicode 문자 포함 파일 수: {unicode_file_count} 개\n"
    
    # if unicode_file_count > 0:
    #     report_content += "\n    대상 파일 목록:\n"
    #     for idx, file in enumerate(unicode_files, 1):
    #         report_content += f"    {idx:2d}. {file}\n"
    #         unicode_count = df[df['FileName'] == file]['HasUnicode_Cnt'].iloc[0]
    #         report_content += f"       Unicode 컬럼 수: {unicode_count} 개\n"
    
    # report_content += "\n" + "=" * 60 + "\n\n"

    # 보고서 저장
    try:
        # REPORT_FILE_Name = f'QDQM_File_Info_Report_{pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")}.txt'
        REPORT_FILE_Name = f'QDQM_File_Info_Report.txt'
        output_dir = os.path.join(CURRENT_DIR_PATH, 'report')
        os.makedirs(output_dir, exist_ok=True)
        
        report_file = os.path.join(output_dir, REPORT_FILE_Name)
        
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(report_content)
            
        st.success(f"보고서가 저장되었습니다: {report_file}")
        
        # 다운로드 버튼 추가
        with open(report_file, 'r', encoding='utf-8') as f:
            st.download_button(
                label="보고서 다운로드",
                data=f.read(),
                file_name=os.path.basename(report_file),
                mime="text/plain"
            )
            
    except Exception as e:
        st.error(f"보고서 저장 중 오류가 발생했습니다: {str(e)}")


def generate_markdown_report(df, CURRENT_DIR_PATH, QDQM_ver):
    if df.empty:
        st.warning("데이터가 없습니다.")
        return

    # 보고서 내용 생성
    report_content = f"""# Data Quality Management Report

## QDQM Analyzer Report v{QDQM_ver}

- 작성자: QDQM Analyzer
- 작성일시: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}

> 본 보고서는 QDQM Analyzer에 의해 자동 생성된 보고서입니다.  
> 데이터 품질 관리를 위한 자료로 활용하시기 바랍니다.
> 데이터 품질관리에 대한 자세한 내용은 [여기](https://qliksense.tistory.com/176)를 참조하세요.

---

## Tables Information

"""

    # 기본 정보 추가
    summary = master_table_summary(df)
    for key, value in summary.items():
        report_content += f"- **{key}**: {value}\n"

    # 특수 케이스 파일 수 정보
    unicode_file_count = len(df[df['UnicodeCnt'] > 0])
    broken_kor_file_count = len(df[df['BrokenKoreanCnt'] > 0])
    lt_lcl_file_count = len(df[df['BelowLCL'] > 0])
    gt_ucl_file_count = len(df[df['UpperUCL'] > 0])

    report_content += "\n### 특수 케이스 분석\n"
    report_content += f"- Unicode 문자 포함 파일 수: **{unicode_file_count}** 개\n"
    report_content += f"- 한글 미완성 문자 포함 파일 수: **{broken_kor_file_count}** 개\n"
    report_content += f"- Less Than LCL 포함 파일 수: **{lt_lcl_file_count}** 개\n"
    report_content += f"- Greater Than UCL 포함 파일 수: **{gt_ucl_file_count}** 개\n"
    
    report_content += f"\n> LCL, UCL에 관한 자세한 내용은 [여기](https://qliksense.tistory.com/45)를 참조하세요.\n"

    # 파일별 상세 정보
    report_content += "\n## Tables Detail Information\n"
    
    for _, row in df.iterrows():
        report_content += f"\n### {row['FileName']}\n\n"
        
        report_content += "#### 기본 정보\n"
        report_content += f"- Columns #: {row['ColumnCnt']:,}\n"
        report_content += f"- File Size: {row['FileSize']/1000:,.0f} KB\n"
        report_content += f"- Row #: {row['RecordCnt']:,.0f}\n"
        report_content += f"- Sampling #: {row['SamplingRows']:,.0f}\n"
        report_content += f"- Unique Columns: {row['UniqueCnt']}\n"

        report_content += "\n#### Column # by Data Type\n"
        report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
        report_content += "|------|------|------|------|\n"
        report_content += f"| Text Columns | {row['IsText']} | Null Columns | {row['NullCnt']} |\n"
        report_content += f"| Numeric Columns | {row['IsNum']} | Varchar(Num) | {row['NumChar']} |\n"
        report_content += f"| Date Columns | {row['Date']} | Varchar(Date) | {row['DateChar']} |\n"
        report_content += f"| Varchar(Time) | {row['TimeChar']} | TimeStamp Columns | {row['TimeStamp']} |\n"

        report_content += "\n#### Column # by Data Value\n"
        report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
        report_content += "|------|------|------|------|\n"
        report_content += f"| Has Alpha | {row['HasAlpha']} | Has Only Alpha | {row['HasOnlyAlpha']} |\n"
        report_content += f"| Only Alpha&Num | {row['HasOnlyAlphaNum']} | Has Kor | {row['HasKor']} |\n"
        report_content += f"| Has Only Kor | {row['HasOnlyKor']} | Broken Kor | {row['BrokenKoreanCnt']} |\n"
        report_content += f"| Has Special | {row['SpecialCnt']} | Has Unicode | {row['UnicodeCnt']} |\n"
        report_content += f"| Has Chinese | {row['ChineseCnt']} | Has Null | {row['NullCnt']} |\n"
        report_content += f"| Has Blank | {row['HasBlank']} | Has Dash | {row['HasDash']} |\n"
        report_content += f"| Has Dot | {row['HasDot']} | Has Bracket | {row['HasBracket']} |\n"

        report_content += "\n#### Column # by Numeric Value & Statistics\n"
        report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
        report_content += "|------|------|------|------|\n"
        report_content += f"| Has Only Num | {row['HasOnlyNum']} | Has Minus | {row['HasMinus']} |\n"
        report_content += f"| Has Float | {row['HasNum']} | Less Than LCL | {row['BelowLCL']} |\n"
        report_content += f"| Greater Than UCL | {row['UpperUCL']} |\n"

        report_content += "\n#### Column # by Master Matching\n"
        report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
        report_content += "|------|------|------|------|\n"
        report_content += f"| 적합 | {row['적합']} | 양호 | {row['양호']} |\n"
        report_content += f"| 고려 | {row['고려']} | 부적합 | {row['부적합']} |\n"
        report_content += f"| Not Matched | {row['X']} | Matching Ratio | {(row['적합']+row['양호'])/(row['ColumnCnt']):,.0%} |\n"
        
        report_content += "\n> Matching Ratio = (적합 + 양호) / 컬럼수\n"
        report_content += "\tMatching Ratio = (적합 + 양호) / 컬럼수\n"
        report_content += "\n---\n"

    # 보고서 저장
    try:
        # REPORT_FILE_Name = f'QDQM_File_Info_Report_{pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")}.md'
        REPORT_FILE_Name = f'QDQM_File_Info_Report.md'
        output_dir = os.path.join(CURRENT_DIR_PATH, 'report')
        os.makedirs(output_dir, exist_ok=True)
        
        report_file = os.path.join(output_dir, REPORT_FILE_Name)
        
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write(report_content)
            
        st.success(f"마크다운 보고서가 저장되었습니다: {report_file}")
        
        # 다운로드 버튼 추가
        with open(report_file, 'r', encoding='utf-8') as f:
            st.download_button(
                label="마크다운 보고서 다운로드",
                data=f.read(),
                file_name=os.path.basename(report_file),
                mime="text/markdown"
            )
            
    except Exception as e:
        st.error(f"보고서 저장 중 오류가 발생했습니다: {str(e)}")      


def Display_Footer():
    st.markdown("""
    \n
##### 파일의 상세 분석결과는 좌측의 
<div style="
    color: #1f77b4; 
    font-size: 24px; 
    font-weight: bold; 
    display: inline-block;
    background-color: #f0f8ff;
    padding: 4px 8px;
    border-radius: 4px;
    margin: 4px 0;
">Tables Information Detail</div>
<span style="font-size: 24px;"> 메뉴를 선택하세요.</span>
""", unsafe_allow_html=True)          



def generate_word_report(df, CURRENT_DIR_PATH, QDQM_ver):
    if df.empty:
        st.warning("데이터가 없습니다.")
        return
        
    try:
        # Word 문서 생성
        doc = Document()
        
        # 제목 추가
        title = doc.add_heading('QDQM 파일 정보 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성일자 추가
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f'생성일자: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 1. 기본 통계 정보를 표로 작성
        doc.add_heading('1. 기본 통계 정보', level=1)
        summary = master_table_summary(df)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = '항목'
        header_cells[1].text = '값'
        
        for key, value in summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # 간격 추가
        
        # 2. 데이터 상세 분석
        doc.add_heading('2. 데이터 상세 분석', level=1)
        
        # 파일별 상세 정보를 표로 작성
        for _, row in df.iterrows():
            doc.add_heading(f'파일: {row["FileName"]}', level=2)
            
            # 기본 정보 파일
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = '항목'
            table.rows[0].cells[1].text = '값'
            
            # 기본 정보 추가
            basic_info = [
                ('Columns #', f"{row['ColumnCnt']:,}"),
                ('File Size', f"{row['FileSize']/1000:,.0f} KB"),
                ('Row #', f"{row['RecordCnt']:,.0f}"),
                ('Sampling #', f"{row['SamplingRows']:,.0f}"),
                ('Text Columns', f"{row['IsText']:,}"),
                ('Numeric Columns', f"{row['IsNum']:,}"),
                ('Date Columns', f"{row['Date']:,}"),
                ('Unicode Columns', f"{row['UnicodeCnt']:,}"),
                ('Broken Korean', f"{row['BrokenKoreanCnt']:,}"),
                ('Chinese Characters', f"{row['ChineseCnt']:,}"),
                ('Sample Rate', f"{row['SamplingRows']/row['RecordCnt']:.1%}")
            ]
            
            for item, value in basic_info:
                cells = table.add_row().cells
                cells[0].text = item
                cells[1].text = str(value)
            
            doc.add_paragraph()  # 간격 추가
            
            # 구분선 추가
            doc.add_paragraph('─' * 50)
        
        # 3. 특이사항 요약
        doc.add_heading('3. 특이사항 요약', level=1)
        summary_table = doc.add_table(rows=1, cols=2)
        summary_table.style = 'Table Grid'
        summary_table.rows[0].cells[0].text = '항목'
        summary_table.rows[0].cells[1].text = '파일 수'
        
        summary_items = [
            ('Unicode 문자 포함', len(df[df['UnicodeCnt'] > 0])),
            ('한글 깨짐 문자 포함', len(df[df['BrokenKoreanCnt'] > 0])),
            ('한자 포함', len(df[df['ChineseCnt'] > 0])),
            ('NULL 값 포함', len(df[df['NullCnt'] > 0])),
            ('특수문자 포함', len(df[df['SpecialCnt'] > 0]))
        ]
        
        for item, count in summary_items:
            cells = summary_table.add_row().cells
            cells[0].text = item
            cells[1].text = f"{count:,} 개"
        
        # 저장
        report_file = os.path.join(CURRENT_DIR_PATH, 'report', 'QDQM_Tables_Report.docx')
        os.makedirs(os.path.dirname(report_file), exist_ok=True)
        doc.save(report_file)
        
        st.success(f"Word 보고서가 생성되었습니다: {report_file}")
        
        # 다운로드 버튼 추가
        with open(report_file, 'rb') as f:
            st.download_button(
                label="Word 보고서 다운로드",
                data=f.read(),
                file_name="QDQM_Tables_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    except Exception as e:
        st.error(f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

def set_table_width(table):
    """파일의 모든 열 너비를 동일하게 설정하는 헬퍼 함수"""
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(10 / len(row.cells))  # A4 가로 기준 약 9인치를 열 개수로 나눔
            # 셀 내용 가운데 정렬
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
