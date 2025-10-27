import pandas as pd
import os
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.section import WD_ORIENT, WD_ORIENTATION
from datetime import datetime

# Create Main KPIs
def master_File_summary(df):
    file_cnt = len(df['FileName'].unique())
    Column_cnt = df['ColumnCnt'].sum()
    record_Cnt = df['RecordCnt'].sum()/100000
    file_size = df['FileSize'].sum()/1000000
    work_date = df['WorkDate'].unique().max()
    return {
        "Total Files #": f"{file_cnt:,}", 
        "Total Columns #": f"{Column_cnt:,}", # Column_cnt,
        "Record #(백만)": f"{record_Cnt:,.0f}",   # record_Cnt, 
        "File Size(MB)": f"{file_size:,.0f}",  #file_size,
        "Analyzed Date" : work_date
    }

def Display_KPIs(df):
    st.markdown("####  분석 대상 파일(파일)들의 주요 지표 입니다.")   
    # CSS 스타일 정의
    st.markdown("""
        <style>
        div[data-testid="metric-container"] {
            background-color: #FFFFFF;
            border: 1px solid #E0E0E0;
            padding: 1rem;
            border-radius: 10px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        div[data-testid="metric-container"] > label[data-testid="stMetricLabel"] {
            display: flex;
            justify-content: center;
            color: #404040;
            font-weight: 600;
        }
        div[data-testid="metric-container"] > div[data-testid="stMetricValue"] {
            display: flex;
            justify-content: center;
            font-weight: bold;
        }
        </style>
    """, unsafe_allow_html=True)

    # KPI 표시
    summary = master_File_summary(df)
    
    # 각 메트릭에 대한 색상 정의
    metric_colors = {
        "Total Files #": "#1f77b4",      # 파란색
        "Total Columns #": "#2ca02c",      # 초록색
        "Record #(백만)": "#ff7f0e",      # 주황색
        "File Size(MB)": "#d62728",       # 빨간색
        "Working Date": "#9467bd"         # 보라색
    }

    # 메트릭 표시
    cols = st.columns(len(summary))
    for col, (key, value) in zip(cols, summary.items()):
        color = metric_colors.get(key, "#0072B2")  # 기본 색상
        col.markdown(f"""
            <div style="text-align: center; padding: 1.2rem; background-color: #FFFFFF; border-radius: 10px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                <div style="color: {color}; font-size: 2em; font-weight: bold;">{value}</div>
                <div style="color: #404040; font-size: 1em; margin-top: 0.5rem;">{key}</div>
            </div>
        """, unsafe_allow_html=True)

# # -------------------------------------------------------------------------------
# # 보고서 생성 함수
# # -------------------------------------------------------------------------------
# def generate_text_report(df, CURRENT_DIR_PATH, QDQM_ver):
#     if df.empty:
#         st.warning("데이터가 없습니다.")
#         return

#     # 보고서 내용 생성
#     report_content = "\n"  
#     report_content += """
# ══════════════════════════════════════════════════════════════
                                                             
#                Data Quality Management Report                 
                                                              
#                    QDQM Analyzer Report                      
                                                             
# ══════════════════════════════════════════════════════════════
# """
#     report_content += "\n"*3   
#     report_content += f"""
#     \t작성자\t\t: QDQM Analyzer
#     \t작성일시\t\t: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}
#     \t버전\t\t: {QDQM_ver}
#     """
#     report_content += "\n"*3

#     report_content += """
#     \t본 보고서는 QDQM Analyzer에 의해 자동 생성된 보고서입니다.
#     \t데이터 품질 관리를 위한 자료로 활용하시기 바랍니다.
#     """
#     report_content += "\n" + "=" * 60 + "\n\n"
    
#     # 기본 정보 추가
#     summary = master_File_summary(df)
#     report_content += "[ Files Information ]\n\n"
#     for key, value in summary.items():
#         report_content += f"\t{key}\t: {value}\n"

#     unicode_files = df[df['UnicodeCnt'] > 0]['FileName'].tolist()
#     unicode_file_count = len(unicode_files)
#     report_content += f"\n\tUnicode 문자 포함 파일 수: {unicode_file_count} 개\n"
    
#     broken_kor_files = df[df['BrokenKoreanCnt'] > 0]['FileName'].tolist()
#     broken_kor_file_count = len(broken_kor_files)
#     report_content += f"\t한글 미완성 문자 포함 파일 수: {broken_kor_file_count} 개\n"
    
#     lt_lcl_files = df[df['BelowLCL'] > 0]['FileName'].tolist()
#     lt_lcl_file_count = len(lt_lcl_files)
#     report_content += f"\tLCL 보다 작은 값을 갖고있는 파일 수: {lt_lcl_file_count} 개\n"  
    
#     gt_ucl_files = df[df['UpperUCL'] > 0]['FileName'].tolist()
#     gt_ucl_file_count = len(gt_ucl_files)
#     report_content += f"\tUCL 보다 큰 값을 갖고있는 파일 수: {gt_ucl_file_count} 개\n"   
    
      
#     LCLUCL_Site = "https://qliksense.tistory.com/45"
#     report_content += f"\n\tLCL, UCL에 관한 자세한 내용은 여기를 참조하세요: {LCLUCL_Site}\n"
    
#     report_content += "\n" + "=" * 60 + "\n\n"

#     # 파일별 상세 정보
#     report_content += "[ Files Detail Information ]\n"
#     report_content += "\n" + "=" * 60 + "\n\n"
#     for _, row in df.iterrows():
#         report_content += f"[ {row['FileName']} ] File Detail Information \n\n"
#         report_content += f"\tColumns #\t\t: {row['ColumnCnt']}\n"
#         report_content += f"\tFile Size\t\t: {row['FileSize']/1000:,.0f} KB\n"
#         if len(str(int(row['RecordCnt']))) > 6:
#             report_content += f"\tRow #\t\t: {row['RecordCnt']:,.0f}\n"
#         else:
#             report_content += f"\tRow #\t\t: {row['RecordCnt']:,.0f}\n"
#         report_content += f"\tSampling #\t: {row['SamplingRows']:,.0f}\n"
#         report_content += f"\tUnique Columns\t: {row['UniqueCnt']}\n"
#         report_content += "\n\t--- Column # by Data Type ---\n"
#         report_content += f"\tText Columns\t: {row['IsText']}\t\t"
#         report_content += f"\tNull Columns\t: {row['IsNull']}\n"
#         report_content += f"\tNumeric Columns\t: {row['IsNum']}\t\t"
#         report_content += f"\tVarchar(Num)\t: {row['NumChar']}\n"    
#         report_content += f"\tDate Columns\t: {row['Date']}\t\t"
#         report_content += f"\tVarchar(Date)\t: {row['DateChar']}\n"
#         report_content += f"\tVarchar(Time)\t: {row['TimeChar']}\t\t"
#         report_content += f"\tTimeStamp Columns\t: {row['TimeStamp']}\n" 
#         report_content += "\n\t--- Column # by Data Value ---\n"
#         report_content += f"\tHas Alpha     \t: {row['HasAlpha']}\n"
#         report_content += f"\tHas Only Alpha\t: {row['HasOnlyAlpha']}\t\t"
#         report_content += f"\tOnly Alpha&Num\t: {row['HasOnlyAlphaNum']}\n"
#         report_content += f"\tHas Kor    \t: {row['HasKor']}\n"
#         report_content += f"\tHas Only Kor\t: {row['HasOnlyKor']}\t\t"
#         report_content += f"\tBroken Kor\t: {row['BrokenKoreanCnt']}\n"
#         report_content += f"\tHas Special\t: {row['SpecialCnt']}\t\t"
#         report_content += f"\tHas Unicode\t: {row['UnicodeCnt']}\n"
#         report_content += f"\tHas Chinese\t: {row['ChineseCnt']}\t\t"   
#         report_content += f"\tHas Null   \t: {row['NullCnt']}\n"
#         report_content += f"\tHas Blank  \t: {row['HasBlank']}\t\t"
#         report_content += f"\tHas Dash  \t: {row['HasDash']}\n"
#         report_content += f"\tHas Dot    \t: {row['HasDot']}\t\t"
#         report_content += f"\tHas Bracket\t: {row['HasBracket']}\n"
#         # report_content += f"Has At\t: {row['HasAt_Cnt']}\n"
#         report_content += "\n\t--- Column # by Numeric Value & Statistics ---\n"
#         report_content += f"\tHas Only Num\t: {row['HasOnlyNum']}\n"
#         report_content += f"\tHas Minus   \t: {row['HasMinus']}\t\t"
#         # report_content += f"\tHas Float  \t: {row['HasReal']}\n"
#         report_content += f"\tLess Than LCL\t: {row['BelowLCL']}\t\t"
#         report_content += f"\tGreater Than UCL\t: {row['UpperUCL']}\n"

#         report_content += "\n\t--- Column # by Master Matching ---\n"
#         report_content += f"\t적합        \t: {row['적합']}\t\t"
#         report_content += f"\t양호        \t: {row['양호']}\n"
#         report_content += f"\t고려        \t: {row['고려']}\t\t"
#         report_content += f"\t부적합      \t: {row['부적합']}\n"
#         report_content += f"\tNot Matched  \t: {row['X']}\t\t"
#         report_content += f"\tMatching Ratio\t: {(row['적합']+row['양호'])/(row['ColumnCnt']):,.0%}\n"
#         report_content += "\tMatching Ratio = (적합 + 양호) / 컬럼수\n"
#         report_content += "-" * 100 + "\n\n"

#     # # Unicode 문자가 포함된 파일 분석
#     # unicode_files = df[df['HasUnicode_Cnt'] > 0]['FileName'].tolist()
#     # unicode_file_count = len(unicode_files)
    
#     # report_content += "\n[ Unicode 문자 포함 파일 분석 ]\n\n"
#     # report_content += f"    Unicode 문자 포함 파일 수: {unicode_file_count} 개\n"
    
#     # if unicode_file_count > 0:
#     #     report_content += "\n    대상 파일 목록:\n"
#     #     for idx, file in enumerate(unicode_files, 1):
#     #         report_content += f"    {idx:2d}. {file}\n"
#     #         unicode_count = df[df['FileName'] == file]['HasUnicode_Cnt'].iloc[0]
#     #         report_content += f"       Unicode 컬럼 수: {unicode_count} 개\n"
    
#     # report_content += "\n" + "=" * 60 + "\n\n"

#     # 보고서 저장
#     try:
#         # REPORT_FILE_Name = f'QDQM_File_Info_Report_{pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")}.txt'
#         REPORT_FILE_Name = f'QDQM_File_Info_Report.txt'
#         output_dir = os.path.join(CURRENT_DIR_PATH, 'report')
#         os.makedirs(output_dir, exist_ok=True)
        
#         report_file = os.path.join(output_dir, REPORT_FILE_Name)
        
#         with open(report_file, 'w', encoding='utf-8') as f:
#             f.write(report_content)
            
#         st.success(f"보고서가 저장되었습니다: {report_file}")
        
#         # 다운로드 버튼 추가
#         with open(report_file, 'r', encoding='utf-8') as f:
#             st.download_button(
#                 label="보고서 다운로드",
#                 data=f.read(),
#                 file_name=os.path.basename(report_file),
#                 mime="text/plain"
#             )
            
#     except Exception as e:
#         st.error(f"보고서 저장 중 오류가 발생했습니다: {str(e)}")


# def generate_markdown_report(df, CURRENT_DIR_PATH, QDQM_ver):
#     if df.empty:
#         st.warning("데이터가 없습니다.")
#         return

#     # 보고서 내용 생성
#     report_content = f"""# Data Quality Management Report

# ## QDQM Analyzer Report v{QDQM_ver}

# - 작성자: QDQM Analyzer
# - 작성일시: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}

# > 본 보고서는 QDQM Analyzer에 의해 자동 생성된 보고서입니다.  
# > 데이터 품질 관리를 위한 자료로 활용하시기 바랍니다.
# > 데이터 품질관리에 대한 자세한 내용은 [여기](https://qliksense.tistory.com/176)를 참조하세요.

# ---

# ## Files Information

# """

#     # 기본 정보 추가
#     summary = master_File_summary(df)
#     for key, value in summary.items():
#         report_content += f"- **{key}**: {value}\n"

#     # 특수 케이스 파일 수 정보
#     unicode_file_count = len(df[df['UnicodeCnt'] > 0])
#     broken_kor_file_count = len(df[df['BrokenKoreanCnt'] > 0])
#     lt_lcl_file_count = len(df[df['BelowLCL'] > 0])
#     gt_ucl_file_count = len(df[df['UpperUCL'] > 0])

#     report_content += "\n### 특수 케이스 분석\n"
#     report_content += f"- Unicode 문자 포함 파일 수: **{unicode_file_count}** 개\n"
#     report_content += f"- 한글 미완성 문자 포함 파일 수: **{broken_kor_file_count}** 개\n"
#     report_content += f"- Less Than LCL 포함 파일 수: **{lt_lcl_file_count}** 개\n"
#     report_content += f"- Greater Than UCL 포함 파일 수: **{gt_ucl_file_count}** 개\n"
    
#     report_content += f"\n> LCL, UCL에 관한 자세한 내용은 [여기](https://qliksense.tistory.com/45)를 참조하세요.\n"

#     # 파일별 상세 정보
#     report_content += "\n## Files Detail Information\n"
    
#     for _, row in df.iterrows():
#         report_content += f"\n### {row['FileName']}\n\n"
        
#         report_content += "#### 기본 정보\n"
#         report_content += f"- Columns #: {row['ColumnCnt']:,}\n"
#         report_content += f"- File Size: {row['FileSize']/1000:,.0f} KB\n"
#         report_content += f"- Row #: {row['RecordCnt']:,.0f}\n"
#         report_content += f"- Sampling #: {row['SamplingRows']:,.0f}\n"
#         report_content += f"- Unique Columns: {row['UniqueCnt']}\n"

#         report_content += "\n#### Column # by Data Type\n"
#         report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
#         report_content += "|------|------|------|------|\n"
#         report_content += f"| Text Columns | {row['IsText']} | Null Columns | {row['NullCnt']} |\n"
#         report_content += f"| Numeric Columns | {row['IsNum']} | Varchar(Num) | {row['NumChar']} |\n"
#         report_content += f"| Date Columns | {row['Date']} | Varchar(Date) | {row['DateChar']} |\n"
#         report_content += f"| Varchar(Time) | {row['TimeChar']} | TimeStamp Columns | {row['TimeStamp']} |\n"

#         report_content += "\n#### Column # by Data Value\n"
#         report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
#         report_content += "|------|------|------|------|\n"
#         report_content += f"| Has Alpha | {row['HasAlpha']} | Has Only Alpha | {row['HasOnlyAlpha']} |\n"
#         report_content += f"| Only Alpha&Num | {row['HasOnlyAlphaNum']} | Has Kor | {row['HasKor']} |\n"
#         report_content += f"| Has Only Kor | {row['HasOnlyKor']} | Broken Kor | {row['BrokenKoreanCnt']} |\n"
#         report_content += f"| Has Special | {row['SpecialCnt']} | Has Unicode | {row['UnicodeCnt']} |\n"
#         report_content += f"| Has Chinese | {row['ChineseCnt']} | Has Null | {row['NullCnt']} |\n"
#         report_content += f"| Has Blank | {row['HasBlank']} | Has Dash | {row['HasDash']} |\n"
#         report_content += f"| Has Dot | {row['HasDot']} | Has Bracket | {row['HasBracket']} |\n"

#         report_content += "\n#### Column # by Numeric Value & Statistics\n"
#         report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
#         report_content += "|------|------|------|------|\n"
#         report_content += f"| Has Only Num | {row['HasOnlyNum']} | Has Minus | {row['HasMinus']} |\n"
#         report_content += f"| Has Float | {row['HasNum']} | Less Than LCL | {row['BelowLCL']} |\n"
#         report_content += f"| Greater Than UCL | {row['UpperUCL']} |\n"

#         report_content += "\n#### Column # by Master Matching\n"
#         report_content += "| 구분 | 개수 | 구분 | 개수 |\n"
#         report_content += "|------|------|------|------|\n"
#         report_content += f"| 적합 | {row['적합']} | 양호 | {row['양호']} |\n"
#         report_content += f"| 고려 | {row['고려']} | 부적합 | {row['부적합']} |\n"
#         report_content += f"| Not Matched | {row['X']} | Matching Ratio | {(row['적합']+row['양호'])/(row['ColumnCnt']):,.0%} |\n"
        
#         report_content += "\n> Matching Ratio = (적합 + 양호) / 컬럼수\n"
#         report_content += "\tMatching Ratio = (적합 + 양호) / 컬럼수\n"
#         report_content += "\n---\n"

#     # 보고서 저장
#     try:
#         # REPORT_FILE_Name = f'QDQM_File_Info_Report_{pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")}.md'
#         REPORT_FILE_Name = f'QDQM_File_Info_Report.md'
#         output_dir = os.path.join(CURRENT_DIR_PATH, 'report')
#         os.makedirs(output_dir, exist_ok=True)
        
#         report_file = os.path.join(output_dir, REPORT_FILE_Name)
        
#         with open(report_file, 'w', encoding='utf-8') as f:
#             f.write(report_content)
            
#         st.success(f"마크다운 보고서가 저장되었습니다: {report_file}")
        
#         # 다운로드 버튼 추가
#         with open(report_file, 'r', encoding='utf-8') as f:
#             st.download_button(
#                 label="마크다운 보고서 다운로드",
#                 data=f.read(),
#                 file_name=os.path.basename(report_file),
#                 mime="text/markdown"
#             )
            
#     except Exception as e:
#         st.error(f"보고서 저장 중 오류가 발생했습니다: {str(e)}")      


# def Display_Footer():
#     st.markdown("""
#     \n
# ##### 파일의 상세 분석결과는 좌측의 
# <div style="
#     color: #1f77b4; 
#     font-size: 24px; 
#     font-weight: bold; 
#     display: inline-block;
#     background-color: #f0f8ff;
#     padding: 4px 8px;
#     border-radius: 4px;
#     margin: 4px 0;
# ">Files Information Detail</div>
# <span style="font-size: 24px;"> 메뉴를 선택하세요.</span>
# """, unsafe_allow_html=True)          



# def generate_word_report(df, CURRENT_DIR_PATH, QDQM_ver):
#     if df.empty:
#         st.warning("데이터가 없습니다.")
#         return
        
#     try:
#         # Word 문서 생성
#         doc = Document()
        
#         # 제목 추가
#         title = doc.add_heading('QDQM 파일 정보 분석 보고서', 0)
#         title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
#         # 생성일자 추가
#         date_paragraph = doc.add_paragraph()
#         date_run = date_paragraph.add_run(f'생성일자: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
#         date_run.font.size = Pt(10)
#         date_run.font.color.rgb = RGBColor(128, 128, 128)
        
#         # 1. 기본 통계 정보를 표로 작성
#         doc.add_heading('1. 기본 통계 정보', level=1)
#         summary = master_File_summary(df)
#         File = doc.add_table(rows=1, cols=2)
#         File.style = 'Table Grid'
#         header_cells = File.rows[0].cells
#         header_cells[0].text = '항목'
#         header_cells[1].text = '값'
        
#         for key, value in summary.items():
#             row_cells = File.add_row().cells
#             row_cells[0].text = key
#             row_cells[1].text = str(value)
        
#         doc.add_paragraph()  # 간격 추가
        
#         # 2. 데이터 상세 분석
#         doc.add_heading('2. 데이터 상세 분석', level=1)
        
#         # 파일별 상세 정보를 표로 작성
#         for _, row in df.iterrows():
#             doc.add_heading(f'파일: {row["FileName"]}', level=2)
            
#             # 기본 정보 파일
#             File = doc.add_table(rows=1, cols=2)
#             File.style = 'Table Grid'
#             File.rows[0].cells[0].text = '항목'
#             File.rows[0].cells[1].text = '값'
            
#             # 기본 정보 추가
#             basic_info = [
#                 ('Columns #', f"{row['ColumnCnt']:,}"),
#                 ('File Size', f"{row['FileSize']/1000:,.0f} KB"),
#                 ('Row #', f"{row['RecordCnt']:,.0f}"),
#                 ('Sampling #', f"{row['SamplingRows']:,.0f}"),
#                 ('Text Columns', f"{row['IsText']:,}"),
#                 ('Numeric Columns', f"{row['IsNum']:,}"),
#                 ('Date Columns', f"{row['Date']:,}"),
#                 ('Unicode Columns', f"{row['UnicodeCnt']:,}"),
#                 ('Broken Korean', f"{row['BrokenKoreanCnt']:,}"),
#                 ('Chinese Characters', f"{row['ChineseCnt']:,}"),
#                 ('Sample Rate', f"{row['SamplingRows']/row['RecordCnt']:.1%}")
#             ]
            
#             for item, value in basic_info:
#                 cells = File.add_row().cells
#                 cells[0].text = item
#                 cells[1].text = str(value)
            
#             doc.add_paragraph()  # 간격 추가
            
#             # 구분선 추가
#             doc.add_paragraph('─' * 50)
        
#         # 3. 특이사항 요약
#         doc.add_heading('3. 특이사항 요약', level=1)
#         summary_File = doc.add_table(rows=1, cols=2)
#         summary_File.style = 'Table Grid'
#         summary_File.rows[0].cells[0].text = '항목'
#         summary_File.rows[0].cells[1].text = '파일 수'
        
#         summary_items = [
#             ('Unicode 문자 포함', len(df[df['UnicodeCnt'] > 0])),
#             ('한글 깨짐 문자 포함', len(df[df['BrokenKoreanCnt'] > 0])),
#             ('한자 포함', len(df[df['ChineseCnt'] > 0])),
#             ('NULL 값 포함', len(df[df['NullCnt'] > 0])),
#             ('특수문자 포함', len(df[df['SpecialCnt'] > 0]))
#         ]
        
#         for item, count in summary_items:
#             cells = summary_File.add_row().cells
#             cells[0].text = item
#             cells[1].text = f"{count:,} 개"
        
#         # 저장
#         report_file = os.path.join(CURRENT_DIR_PATH, 'report', 'QDQM_Files_Report.docx')
#         os.makedirs(os.path.dirname(report_file), exist_ok=True)
#         doc.save(report_file)
        
#         st.success(f"Word 보고서가 생성되었습니다: {report_file}")
        
#         # 다운로드 버튼 추가
#         with open(report_file, 'rb') as f:
#             st.download_button(
#                 label="Word 보고서 다운로드",
#                 data=f.read(),
#                 file_name="QDQM_Files_Report.docx",
#                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#             )
            
#     except Exception as e:
#         st.error(f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

#---------------------------------------------------------------------------------------------------

def set_File_width(File):
    """파일의 모든 열 너비를 동일하게 설정하는 헬퍼 함수"""
    for row in File.rows:
        for cell in row.cells:
            cell.width = Inches(10 / len(row.cells))  # A4 가로 기준 약 9인치를 열 개수로 나눔
            # 셀 내용 가운데 정렬
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return 

def generate_data_quality_report(df, df_Column, CURRENT_DIR_PATH, QDQM_ver):
    if df.empty:
        st.warning("데이터가 없습니다.")
        return
        
    try:
        doc = Document()
        
        # 페이지 여백 설정 (1.5 inches = 1.5 * 914400 twips)
        sections = doc.sections
        for section in sections:
            # Change from inches to centimeters (1 cm = 914400 / 2.54 twips)
            section.left_margin = int(1.5 * 914400 / 2.54)   # 1.5 cm
            section.right_margin = int(1.5 * 914400 / 2.54)  # 1.5 cm
            section.top_margin = int(2 * 914400 / 2.54)   # 2 cm
            section.bottom_margin = int(2 * 914400 / 2.54)  # 2 cm

            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = int(297 * 914400 / 25.4)    # A4 너비 (297mm)
            section.page_height = int(210 * 914400 / 25.4)   # A4 높이 (210mm)
        
            section = doc.sections[0]
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run("QDQM Report")
            run.font.size = Pt(8)
            run.font.name = 'Arial'

        # 데이터 타입 변환
        df['UnicodeCnt'] = df['UnicodeCnt'].fillna(0).astype(int)
        df['BrokenKoreanCnt'] = df['BrokenKoreanCnt'].fillna(0).astype(int)
        df['BelowLCL'] = df['BelowLCL'].fillna(0).astype(int)
        df['UpperUCL'] = df['UpperUCL'].fillna(0).astype(int)
        df['Match_Check'] = df['Match_Check'].fillna(0).astype(int)
        df['Date_Check'] = df['Date_Check'].fillna(0).astype(int)
        df['Len_Check'] = df['Len_Check'].fillna(0).astype(int)
        df['ChineseCnt'] = df['ChineseCnt'].fillna(0).astype(int)

        # 제목 추가
        title = doc.add_heading('데이터 품질 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 생성일자 추가
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f'생성일자: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # 1. 전체 통계 정보
        doc.add_heading('1. 전체 통계 정보', level=1)
        total_Files = len(df['FileName'].unique())
        total_Columns = df['ColumnCnt'].sum()
        total_records = df['RecordCnt'].sum()
        total_filesize = df['FileSize'].sum()
        
        # 기본 통계 파일 생성
        File = doc.add_table(rows=1, cols=2)
        File.style = 'Table Grid'
        header_cells = File.rows[0].cells
        header_cells[0].text = '구   분'
        header_cells[1].text = '값'
        
        stats = [
            ('총 파일 수', f"{total_Files:,}개"),
            ('총 컬럼 수', f"{total_Columns:,}개"),
            ('총 레코드 수', f"{total_records/1000:,.0f}천건"),
            ('총 파일 크기', f"{total_filesize/1000000:,.0f} MB"),
        ]
        
        for key, value in stats:
            row_cells = File.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
        
        set_File_width(File)
        
        # 2. 데이터 품질 지표
        doc.add_heading('2. 데이터 품질 지표', level=1)
        # 파일 헤더 생성
        File = doc.add_table(rows=1, cols=4)
        File.style = 'Table Grid'
        header_cells = File.rows[0].cells
        header_cells[0].text = '점검 항목'
        header_cells[1].text = '파일 수'
        header_cells[2].text = '파일 비율(%)'
        header_cells[3].text = '컬럼 수'
        
        # 점검 항목별 통계
        quality_checks = {
            '유니코드': ('UnicodeCnt', df['UnicodeCnt'] > 0),
            '미완성한글': ('BrokenKoreanCnt', df['BrokenKoreanCnt'] > 0),
            '하한선미만': ('BelowLCL', df['BelowLCL'] > 0),
            '상한선초과': ('UpperUCL', df['UpperUCL'] > 0),
            '참조무결성': ('Match_Check', df['Match_Check'] > 0),
            '일자점검': ('Date_Check', df['Date_Check'] > 0),
            '길이점검': ('Len_Check', df['Len_Check'] > 0),
            '한자문자': ('ChineseCnt', df['ChineseCnt'] > 0)
        }
        
        for check_name, (count_col, condition) in quality_checks.items():
            affected_Files = len(df[condition])
            affected_Columns = df[count_col].sum()
            File_percentage = (affected_Files / total_Files * 100) if total_Files > 0 else 0
            row_cells = File.add_row().cells
            row_cells[0].text = check_name
            row_cells[1].text = f"{affected_Files:,}개"
            row_cells[2].text = f"{File_percentage:.1f}%"
            row_cells[3].text = f"{affected_Columns:,}개"
        
        set_File_width(File)
        
        # 3. 주의 대상 파일 목록
        doc.add_heading('3. 파일 목록', level=1)
        warning_condition = ((df['UnicodeCnt'] > 0) | 
                           (df['BrokenKoreanCnt'] > 0) |
                           (df['BelowLCL'] > 0) |
                           (df['UpperUCL'] > 0) |
                           (df['Match_Check'] > 0) |
                           (df['Date_Check'] > 0) |
                           (df['Len_Check'] > 0))
        
        warning_Files = df[warning_condition]
        
        if not warning_Files.empty:
            File = doc.add_table(rows=1, cols=10)
            File.style = 'Table Grid'
            
            # 헤더 설정
            headers = ['파일명', '컬럼수', '유니코드', '미완성한글', '하한선미만', '상한선이상', 
                '참조무결성', '일자점검', '길이점검', '점검']
            for i, header in enumerate(headers):
                File.rows[0].cells[i].text = header
            
            # 데이터 추가
            for _, row in warning_Files.iterrows():
                row_cells = File.add_row().cells
                row_cells[0].text = str(row['FileName'])
                row_cells[1].text = f"{int(row['ColumnCnt']):,}"
                row_cells[2].text = f"⚠️({int(row['UnicodeCnt'])})" if row['UnicodeCnt'] > 0 else ''
                row_cells[3].text = f"⚠️({int(row['BrokenKoreanCnt'])})" if row['BrokenKoreanCnt'] > 0 else ''
                row_cells[4].text = f"⚠️({int(row['BelowLCL'])})" if row['BelowLCL'] > 0 else ''
                row_cells[5].text = f"⚠️({int(row['UpperUCL'])})" if row['UpperUCL'] > 0 else ''
                row_cells[6].text = f"⚠️({int(row['Match_Check'])})" if row['Match_Check'] > 0 else ''
                row_cells[7].text = f"⚠️({int(row['Date_Check'])})" if row['Date_Check'] > 0 else ''
                row_cells[8].text = f"⚠️({int(row['Len_Check'])})" if row['Len_Check'] > 0 else ''
                row_cells[9].text = '🚨' if (row['Match_Check'] > 0 or 
                                            row['Date_Check'] > 0 or 
                                            # row['Len_Check'] > 0 or 
                                            row['BrokenKoreanCnt'] > 0 or 
                                            row['UnicodeCnt'] > 0 or 
                                            row['BelowLCL'] > 0 or 
                                            row['UpperUCL'] > 0) else ''
            
            set_File_width(File)
        else:
            doc.add_paragraph('점검 대상 파일이 없습니다.')

        doc.add_page_break()

        # 6. 컬럼별 상세 점검 리스트
        doc.add_heading('4. 컬럼별 상세 리스트', level=1)
        date_run.font.size = Pt(10)

        #----------------------------------------------------------------------------
        # FileName으로 그룹화하여 순환
        page_cnt = 0
        for file in df_Column['FileName'].unique():
            page_cnt += 1
            # 각 파일에 대한 데이터 필터링
            file_data = df_Column[df_Column['FileName'] == file]  
            doc.add_heading(f'4.{page_cnt} [ {file} ] 점검 항목별 컬럼수', level=2)

            # Filter data for the specific file
            file_data = df_Column[df_Column['FileName'] == file].copy()

            file_data['점검'] = ''
            file_data.loc[(file_data['UnicodeCnt'] > 0) 
                | (file_data['BrokenKoreanCnt'] > 0)
                | (file_data['BelowLCL'] > 0)
                | (file_data['UpperUCL'] > 0)
                | (file_data['Match_Check'] > 0)
                | (file_data['Date_Check'] > 0)
                # | (file_data['Len_Check'] > 0)
                , '점검'] = '🚨'  # 전체 점검

            file_data['Null(%)'] = file_data['Null(%)'].fillna(0).astype(float)
            file_data['Match_Total'] = file_data['Match_Total'].fillna(0).astype(int)
            file_data['Match_Good'] = file_data['Match_Good'].fillna(0).astype(int)
            file_data['Match_Check'] = file_data['Match_Check'].fillna(0).astype(int)
            file_data['Date_Check'] = file_data['Date_Check'].fillna(0).astype(int)
            file_data['Len_Check'] = file_data['Len_Check'].fillna(0).astype(int)
            file_data['UniqueCnt'] = file_data['UniqueCnt'].fillna(0).astype(int)
            file_data['NullCnt'] = file_data['NullCnt'].fillna(0).astype(int)
            file_data['LenMax'] = file_data['LenMax'].fillna(0).astype(int)
            file_data['ChineseCnt'] = file_data['ChineseCnt'].fillna(0).astype(int)
            file_data['SpecialCnt'] = file_data['SpecialCnt'].fillna(0).astype(int)
            file_data['ValueCnt'] = file_data['ValueCnt'].fillna(0).astype(int)
            file_data['Matched'] = file_data['Matched'].fillna(0).astype(int)


            # Create a summary of integrity issues
            integrity_summary = {
                '컬럼 / 행 수': [int(file_data['No'].max()), 
                            f"{int(file_data['RecordCnt'].max()):,}",
                         '총 컬럼 및 행 수'],
                'Unique 컬럼': [
                    len(file_data[file_data['Unique(%)'] == 100]) if not file_data.empty else '',
                    ' ',
                    '데이터가 유니크한 값을 갖는 컬럼 수' 
                ],        
                'NULL컬럼': [
                    len(file_data[file_data['Null(%)'] == 100]) if not file_data.empty else '',
                    f"{len(file_data[file_data['Null(%)'] == 100])/int(file_data['No'].max()):.1%}" if not file_data.empty else '',
                    '전체 값이 NULL인 컬럼 수'
                ],
                'NULL > 50%': [
                    len(file_data[(file_data['Null(%)'] > 50) & (file_data['Null(%)'] < 100)]) if not file_data.empty else '',
                    f"{len(file_data[(file_data['Null(%)'] > 50) & (file_data['Null(%)'] < 100)])/int(file_data['No'].max()):.1%}" if not file_data.empty else '',
                    'NULL 값이 50% 초과인 컬럼 수 (Null 컬럼는 제외)'
                ],
                '점검컬럼': [
                    len(file_data[file_data['점검'] == '🚨']),
                    f"{len(file_data[file_data['점검'] == '🚨'])/int(file_data['No'].max()):.1%}" if not file_data.empty else '',
                    '유니코드, 미완성한글, 하한선미만, 상한선초과, 참조무결성, 일자점검, 길이점검 중 점검대상'
                ],
                '유니코드': [
                    len(file_data[file_data['UnicodeCnt'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['UnicodeCnt'].sum())}",
                    '유니코드를 포함한 컬럼 수 및 행 수'
                ],
                '미완성한글': [
                    len(file_data[file_data['BrokenKoreanCnt'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['BrokenKoreanCnt'].sum())}",
                    '미완성한글을 포함한 컬럼 수 및 행 수'
                ],
                '하한선미만': [
                    len(file_data[file_data['BelowLCL'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['BelowLCL'].sum())}",
                    '하한선미만 값을 포함한 컬럼 수'
                ],
                '상한선이상': [
                    len(file_data[file_data['UpperUCL'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['UpperUCL'].sum())}",
                    '상한선초과 값을 포함한 컬럼 수'
                ],
                '마스터참조': [
                    len(file_data[file_data['Match_Total'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['Match_Total'].sum()) / int(file_data['No'].max()):.1%}",
                    '마스터 코드를 참조하는 컬럼 수 및 참조율(마스터참조/컬럼수)' 
                ],
                '참조적합도': [
                    len(file_data[file_data['Match_Good'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['Match_Good'].sum()) / int(file_data['Match_Total'].sum()):.1%}",
                    '마스터 코드 참조가 적합한 컬럼 수 및 적합율(마스터참조적합컬럼/마스터참조)'
                ],
                '참조무결성': [
                    len(file_data[file_data['Match_Check'] > 0]) if not file_data.empty else '',
                    f"{int(file_data[file_data['Match_Check'] > 0]['ValueCnt'].sum() - file_data[file_data['Match_Check'] > 0]['Matched'].sum()):,}",
                    '참조무결성 점검 대상 컬럼 및 행 수(검사할 컬럼의 값이 마스터에 없는 컬럼 수)'
                ],
                '일자점검': [
                    len(file_data[file_data['Date_Check'] > 0]) if not file_data.empty else '',
                    f"{int(file_data[file_data['Date_Check'] > 0]['ValueCnt'].sum() - file_data[file_data['Date_Check'] > 0]['Matched'].sum()):,}",
                    '일자(문자)형으로 지정된 컬럼에서 값이 점검 대상인 컬럼 수 및 행 수 (ex:20240231)'
                ],
                '길이점검': [
                    int(file_data['Len_Check'].sum() if not file_data.empty else 0),
                    '',
                    '길이점검 대상 컬럼 수 (최대 길이 > 길이 최빈값 * 2) '
                ],
                '한자문자': [
                    len(file_data[file_data['ChineseCnt'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['ChineseCnt'].sum()):,}",
                    '한자 문자를 포함한 컬럼 수 및 행 수'
                ],
                '특수문자': [
                    len(file_data[file_data['SpecialCnt'] > 0]) if not file_data.empty else '',
                    f"{int(file_data['SpecialCnt'].sum()):,}",
                    '특수문자를 포함한 컬럼 수 및 행 수'
                ],
                '숫자 & Null 포함': [
                    len(file_data[(file_data['NullCnt'] > 0) & (file_data['DataType'] != 'object')]) if not file_data.empty else '',
                    '', 
                    '데이터 타입이 숫자이면서 NULL이 있는 컬럼 수'
                ],
                '단일값컬럼': [
                    len(file_data[file_data['UniqueCnt'] == 1]) if not file_data.empty else '',
                    '', 
                    '단일값만 갖고있는 컬럼 수'
                ],
                '최대길이 > 200': [
                    len(file_data[file_data['LenMax'] > 200]) if not file_data.empty else '',
                    '', 
                    '데이터의 길이가 200 Byte 이상인 컬럼 수'
                ],
            }

            # 데이터프레임 생성 및 컬럼명 설정
            summary_df = pd.DataFrame(integrity_summary).transpose()
            summary_df.columns = ['컬럼 수', '행 수 or (%)', '설명']  # 컬럼명 직접 지정

            # Word 문서에 파일 추가 (하나의 파일만 생성)
            File = doc.add_table(rows=len(integrity_summary)+1, cols=4)
            File.style = 'Table Grid'

            # 헤더 추가
            headers = ['점검 항목', '컬럼 수', '행 수 or (%)', '설명']
            for i, header in enumerate(headers):
                cell = File.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 데이터 추가
            for idx, (key, values) in enumerate(integrity_summary.items(), start=1):
                cells = File.rows[idx].cells
                cells[0].text = key
                cells[1].text = str(values[0])
                cells[2].text = str(values[1])
                cells[3].text = str(values[2])
                
                # 셀 정렬
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 파일 생성 후 열 너비 조정
            for i, column in enumerate(File.columns):
                for cell in column.cells:
                    # 열 너비 설정
                    if i == 0:  # 첫 번째 열
                        cell.width = Inches(1.5)
                    elif i == 3:  # 네 번째 열
                        cell.width = Inches(6.5)
                    else:  # 나머지 열
                        cell.width = Inches(1)
                        
                    # 셀의 단락 정렬 설정
                    for paragraph in cell.paragraphs:
                        if i == 0 or i == 3:  # 첫 번째와 네 번째 열
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:  # 나머지 열
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


            # 파일 스타일링
            for row in File.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            doc.add_page_break()
            doc.add_heading(f'4.{page_cnt} [ {file} ] 파일 컬럼별 상세리스트', level=2)
            # 상세 점검 파일 생성
            File = doc.add_table(rows=1, cols=10)  # 헤더 설정한 열 수 만큼 지정 
            File.style = 'Table Grid'
            
            # 헤더 설정
            headers = ['컬럼명', '데이터건수', '유니코드', '미완성한글', '하한선미만', '상한선초과'
                       , '참조무결성', '일자점검', '길이점검', '마스터']
            try:
                for i, header in enumerate(headers):
                    cell = File.rows[0].cells[i]
                    try:
                        cell.text = header
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(9)  
                        cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0 
                                                      else WD_ALIGN_PARAGRAPH.CENTER)
                    except IndexError:
                        st.error(f"파일 헤더 '{header}' 설정 중 오류가 발생했습니다.")
            except Exception as e:
                st.error(f"파일 생성 중 오류가 발생했습니다: {str(e)}")

            # ... existing code ...

            # 파일 생성 후 열 너비 조정
            for i, column in enumerate(File.columns):
                for cell in column.cells:
                    # 첫 번째 열은 2배 너비, 나머지 열은 동일 너비
                    if i == 0:
                        cell.width = Inches(1.5)  # 첫 번째 열 2배 너비
                    else:
                        cell.width = Inches(0.5)  # 나머지 열 기본 너비
            
            # 데이터 행 추가
            for _, row in file_data.iterrows():
                row_cells = File.add_row().cells
                
                # 각 열에 데이터 추가
                row_cells[0].text = str(row['ColumnName'])
                row_cells[1].text = f"{int(row['ValueCnt']):,}"
                row_cells[2].text = f"⚠️({int(row['UnicodeCnt'])})" if row['UnicodeCnt'] > 0 else ''
                row_cells[3].text = f"⚠️({int(row['BrokenKoreanCnt'])})" if row['BrokenKoreanCnt'] > 0 else ''
                row_cells[4].text = f"⚠️({int(row['BelowLCL'])})" if row['BelowLCL'] > 0 else ''
                row_cells[5].text = f"⚠️({int(row['UpperUCL'])})" if row['UpperUCL'] > 0 else ''
                row_cells[6].text = f"⚠️({int(row['ValueCnt']-row['Matched'])})" if row['Match_Check'] > 0 else ''
                row_cells[7].text = f"⚠️({int(row['ValueCnt']-row['Matched'])})" if row['Date_Check'] > 0 else ''
                row_cells[8].text = f"⚠️" if row['Len_Check'] > 0 else ''
                row_cells[9].text = str(row['Master']) if pd.notna(row['Master']) else 'None'

                # 셀 가운데 정렬
                for idx, cell in enumerate(row_cells):
                    paragraph = cell.paragraphs[0]
                    # 기존 텍스트를 저장
                    text = paragraph.text
                    # 기존 runs 삭제
                    for run in paragraph.runs:
                        run._element.getparent().remove(run._element)
                    # 새로운 run 생성 및 스타일 적용
                    run = paragraph.add_run(text)
                    run.font.size = Pt(9)
                    
                    # 정렬 설정
                    if idx == 0:  # 첫 번째 열
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:  # 나머지 열
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    idx += 1

           
            # 파일 간 구분을 위한 페이지 나누기 추가
            if file != df_Column['FileName'].unique()[-1]:  # 마지막 파일이 아닌 경우에만 페이지 나누기
                doc.add_page_break()

        # 저장
        report_file = os.path.join(CURRENT_DIR_PATH, 'report', 'QDQM_DataQuality_Report.docx')
        os.makedirs(os.path.dirname(report_file), exist_ok=True)
        doc.save(report_file)
            
        st.success(f"Word 보고서가 생성되었습니다: {report_file}")
        
        # 다운로드 버튼 추가
        with open(report_file, 'rb') as f:
            st.download_button(
                label="데이터 품질 보고서 다운로드",
                data=f.read(),
                file_name="QDQM_DataQuality_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    except Exception as e:
        st.error(f"보고서 생성 중 오류가 발생했습니다: {str(e)}")

def generate_Column_integrity_report(df_Column, CURRENT_DIR_PATH, QDQM_ver):
    doc = Document()
    
    # 문서를 가로 방향으로 설정
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Mm(297)  # A4 가로 길이
    section.page_height = Mm(210)  # A4 세로 길이
    
    # 제목 추가
    doc.add_heading('컬럼별 상세 점검 리스트', level=1)
    
    # 파일 생성
    File = doc.add_table(rows=1, cols=9)
    File.style = 'Table Grid'
    
    # 헤더 설정
    headers = ['파일명', '컬럼명', '유니코드', '미완성한글', '하한선미만', '상한선초과', '마스터점검', '일자점검', '길이점검']
    for i, header in enumerate(headers):
        cell = File.rows[0].cells[i]
        cell.text = header
        # 헤더 스타일 설정
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 데이터 행 추가
    for _, row in df_Column.iterrows():
        row_cells = File.add_row().cells
        
        # 각 열에 데이터 추가
        row_cells[0].text = str(row['FileName'])
        row_cells[1].text = str(row['ColumnName'])
        row_cells[2].text = f"⚠️({int(row['UnicodeCnt'])})" if row['UnicodeCnt'] > 0 else ''
        row_cells[3].text = f"⚠️({int(row['BrokenKoreanCnt'])})" if row['BrokenKoreanCnt'] > 0 else ''
        row_cells[4].text = f"⚠️({int(row['BelowLCL'])})" if row['BelowLCL'] > 0 else ''
        row_cells[5].text = f"⚠️({int(row['UpperUCL'])})" if row['UpperUCL'] > 0 else ''
        row_cells[6].text = f"⚠️({int(row['Match_Check'])})" if row['Match_Check'] > 0 else ''
        row_cells[7].text = f"⚠️({int(row['Date_Check'])})" if row['Date_Check'] > 0 else ''
        row_cells[8].text = f"⚠️({int(row['Len_Check'])})" if row['Len_Check'] > 0 else ''
        
        # 셀 가운데 정렬
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 파일 너비 조정
    File.allow_autofit = True
    for column in File.columns:
        for cell in column.cells:
            cell.width = Inches(1.0)
    
    # 문서 저장
        # 저장
    report_file = os.path.join(CURRENT_DIR_PATH, 'report', 'QDQM_DataQuality_Report_Detail.docx')
    os.makedirs(os.path.dirname(report_file), exist_ok=True)
    doc.save(report_file)
    
    st.success(f"Word 보고서가 생성되었습니다: {report_file}")
    
    # 다운로드 버튼 추가
    with open(report_file, 'rb') as f:
        st.download_button(
            label="데이터 품질 상세 보고서 다운로드",
            data=f.read(),
            file_name="QDQM_DataQuality_Report_Detail.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )    


